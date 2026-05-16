"""Generate the Slovak technical report as a .docx file.

Pulls headline figures from results_polish/, results_slovak/, results_taiwan/
under project/. Embeds five PNG figures and renders ~4 pages of text in plain
academic Slovak.
"""
from __future__ import annotations

from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
PROJECT = ROOT / "project"
OUT_PATH = ROOT / "Hodnotenie_TabFM_pri_chybajucich_datach.docx"

VIZ = {
    "polish": PROJECT / "results_polish" / "visualizations",
    "slovak": PROJECT / "results_slovak" / "visualizations",
    "taiwan": PROJECT / "results_taiwan" / "visualizations",
}

SPLIT_DIR = ROOT / "_split_figs"
SUPTITLE_CROP_PX = 65


def split_figure(src: Path) -> tuple[Path, Path]:
    """Crop the suptitle and split the figure horizontally into two halves.

    Returns (left_half_path, right_half_path).
    """
    SPLIT_DIR.mkdir(exist_ok=True)
    img = Image.open(src)
    w, h = img.size
    cropped = img.crop((0, SUPTITLE_CROP_PX, w, h))
    cw, ch = cropped.size
    mid = cw // 2
    left = cropped.crop((0, 0, mid, ch))
    right = cropped.crop((mid, 0, cw, ch))
    left_path = SPLIT_DIR / (src.stem + "_L.png")
    right_path = SPLIT_DIR / (src.stem + "_R.png")
    left.save(left_path)
    right.save(right_path)
    return left_path, right_path


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            elem = OxmlElement(f"w:{edge}")
            for k, v in kwargs[edge].items():
                elem.set(qn(f"w:{k}"), v)
            tc_borders.append(elem)
    tc_pr.append(tc_borders)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return h


def add_paragraph(doc, text, *, bold=False, italic=False, size=11, align=None,
                  space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_figure(doc, path: Path, caption: str, width_cm: float = 14.5):
    if not path.exists():
        add_paragraph(doc, f"[Obrázok chýba: {path.name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap.paragraph_format.space_after = Pt(8)


def add_split_figure(doc, src: Path, caption_left: str, caption_right: str,
                     width_cm: float = 14.5):
    """Split a wide multi-subplot figure into two halves and embed both."""
    if not src.exists():
        add_paragraph(doc, f"[Obrázok chýba: {src.name}]", italic=True)
        return
    left, right = split_figure(src)
    add_figure(doc, left, caption_left, width_cm=width_cm)
    add_figure(doc, right, caption_right, width_cm=width_cm)


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(2)


def make_table(doc, header, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        cells = t.rows[r_idx].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths_cm:
        for col_idx, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[col_idx].width = Cm(w)
    return t


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Hodnotenie výkonnosti pretrained tabular foundation modelov "
        "pri práci s chýbajúcimi dátami"
    )
    run.bold = True
    run.font.size = Pt(15)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        "Technická správa — semestrálny projekt"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(11)

    add_paragraph(
        doc,
        "Abstrakt — Práca experimentálne porovnáva pretrained tabular "
        "foundation modely (TabPFN, TabICL) s klasickými ML algoritmami "
        "(vrátane CatBoost ako NaN-aware baseline) na troch reálnych "
        "bankruptcy datasetoch (Polish 1-year, Slovak Manufacture 13, "
        "Taiwan Bankruptcy). Skúmame, ako sa výkonnosť modelov mení pri "
        "rôznych mechanizmoch chýbajúcich hodnôt (MCAR, MAR, MNAR) a pri "
        "rastúcom podiele missing dát (5–40 %). Hodnotené sú metriky "
        "vhodné pre nevyváženú klasifikáciu (balanced accuracy, macro F1, "
        "PR-AUC, recall minoritnej triedy). Výsledky ukazujú, že foundation "
        "modely sú v priemere lepšie a stabilnejšie ako klasické modely, "
        "pričom najlepšou stratégiou na všetkých troch datasetoch je "
        "implicitné spracovanie NaN hodnôt modelom (none).",
        italic=True, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "1. Úvod a motivácia", level=1)
    add_paragraph(
        doc,
        "Pretrained tabular foundation modely ako TabPFN [1] a TabICL [2] "
        "predstavujú novú triedu modelov, ktoré sa pretrénovali na veľkom "
        "množstve syntetických tabuľkových úloh a na nový dataset sa "
        "aplikujú bez ďalšieho trénovania alebo s minimálnym fine-tuningom. "
        "V reálnych aplikáciách, najmä v oblasti finančnej predikcie, sú "
        "však tabuľkové dáta často neúplné. Cieľom tejto práce je "
        "experimentálne posúdiť, ako tieto pretrénované modely zvládajú "
        "rôzne typy a úrovne chýbajúcich hodnôt v porovnaní s klasickými "
        "algoritmami strojového učenia.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Hlavnou témou nie je riešenie nevyváženosti tried, ale robustnosť "
        "modelov voči chýbajúcim dátam. Nevyváženosť je však dôležitou "
        "vlastnosťou použitých bankruptcy datasetov [3], a preto sú "
        "metriky volené tak, aby zohľadňovali výrazne nerovnomerné "
        "rozdelenie tried.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "2. Dáta", level=1)
    add_paragraph(
        doc,
        "Experimenty boli realizované na troch reálnych datasetoch z "
        "oblasti predikcie bankrotu, vybraných z prehľadovej štúdie [3]. "
        "Všetky tri datasety obsahujú finančné ukazovatele firiem a "
        "binárnu cieľovú premennú, kde trieda 1 reprezentuje bankrotujúce "
        "firmy. Vstupy boli predspracované, pričom Taiwan je prakticky bez "
        "natívnych NaN hodnôt, zatiaľ čo Polish a Slovak už obsahujú menší "
        "podiel chýbajúcich hodnôt. Nad týmto základom boli následne "
        "kontrolovane simulované scenáre missingness. Charakteristiky datasetov sú zhrnuté v "
        "Tabuľke 1.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    make_table(
        doc,
        ["Dataset", "Počet záznamov", "Trieda 0", "Trieda 1", "Minorita (%)"],
        [
            ["Polish Companies 1-year", "7027", "6756", "271", "3,86"],
            ["Taiwan Bankruptcy", "6819", "6599", "220", "3,23"],
            ["Slovak Manufacture 13", "4107", "4077", "30", "0,73"],
        ],
        widths_cm=[5.0, 3.2, 2.4, 2.4, 2.6],
    )
    add_paragraph(
        doc,
        "Dataset Slovak Manufacture 13 je extrémne náročný — po "
        "stratifikovanom rozdelení 80/20 zostáva v tréningovej množine iba "
        "24 a v testovacej iba 6 bankrotujúcich prípadov. Tento dataset "
        "umožňuje hodnotiť správanie modelov pri veľmi malých minoritných "
        "triedach.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Dôležitý detail pre interpretáciu experimentu je počiatočný stav "
        "chýbajúcich hodnôt ešte pred simuláciou MCAR/MAR/MNAR. V "
        "spracovaných vstupoch má Polish 1-year 5 835 chýbajúcich buniek "
        "(1,28 %), Slovak Manufacture 13 má 12 928 chýbajúcich buniek "
        "(4,84 %) a Taiwan Bankruptcy neobsahuje žiadne natívne NaN "
        "hodnoty (0,00 %). Simulované missingness scenáre teda navyšujú "
        "už existujúce chýbanie pri Polish/Slovak, kým pri Taiwan "
        "vychádzajú z úplne kompletných dát.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "3. Metodológia", level=1)
    add_heading(doc, "3.1 Experimentálny návrh", level=2)
    add_paragraph(
        doc,
        "Pre každý dataset bolo vytvorené stratifikované rozdelenie "
        "train/test v pomere 80/20 s pevným random_state = 42. Rovnaké "
        "splity boli použité pre všetky modely a všetky scenáre, aby bolo "
        "porovnanie objektívne. Chýbajúce hodnoty boli vkladané iba do "
        "tréningovej množiny; testovacia množina zostala kompletná, aby sa "
        "dal merať vplyv neúplných tréningových dát na generalizáciu.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(doc, "Použité boli nasledujúce mechanizmy missingness:")
    add_bullets(doc, [
        "MCAR — Missing Completely At Random; hodnoty sa vynechávajú "
        "úplne náhodne, nezávisle od dát.",
        "MAR — Missing At Random; pravdepodobnosť chýbania závisí od "
        "inej, pozorovanej premennej.",
        "MNAR — Missing Not At Random; pravdepodobnosť chýbania závisí "
        "od samotnej hodnoty premennej.",
    ])
    add_paragraph(
        doc,
        "Pre každý mechanizmus boli simulované úrovne 5 %, 10 %, 15 %, "
        "20 %, 30 % a 40 % chýbajúcich hodnôt. Spoľahlivosť injekcie bola "
        "verifikovaná porovnaním cieľovej a skutočnej miery chýbania "
        "(tolerancia 0,01).",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "3.2 Metódy spracovania chýbajúcich hodnôt", level=2)
    add_paragraph(
        doc,
        "Spôsob, akým naložíme s chýbajúcimi hodnotami, môže výrazne "
        "ovplyvniť kvalitu predikcií. V experimente sme preto porovnali "
        "šesť rôznych stratégií, od najjednoduchších štatistických "
        "imputácií až po implicitné spracovanie NaN priamo modelom. "
        "Všetky imputery boli fitované výlučne na tréningovej množine a "
        "potom aplikované aj na testovaciu množinu, aby sa zabránilo "
        "data leakage zo štatistík testovacích dát:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_bullets(doc, [
        "mean — jednoduchá imputácia priemerom: chýbajúca hodnota v "
        "stĺpci sa nahradí priemerom tohto stĺpca z trénovacích dát. "
        "Rýchla, ale ignoruje vzťahy medzi premennými.",

        "median — jednoduchá imputácia mediánom: rovnaká myšlienka ako "
        "mean, len odolnejšia voči odľahlým hodnotám, čo je vhodné pre "
        "finančné ukazovatele s asymetrickým rozdelením.",

        "kNN — kNN imputácia (KNNImputer, n_neighbors = 5): chýbajúca "
        "hodnota sa nahradí váženým priemerom hodnôt 5 najpodobnejších "
        "tréningových príkladov. Zachytáva lokálnu štruktúru dát, ale "
        "je výpočtovo náročnejšia a citlivá na škálovanie príznakov.",

        "MICE — Multiple Imputation by Chained Equations (scikit-learn "
        "IterativeImputer): každý stĺpec s NaN sa modeluje ako funkcia "
        "ostatných stĺpcov a hodnoty sa iteratívne dopĺňajú regresným "
        "modelom. Lepšie zachytáva multivariantné závislosti ako mean "
        "alebo median.",

        "mice_indicator — MICE s pridanými indikátormi chýbania: ku "
        "každému imputovanému stĺpcu sa pridá nový binárny stĺpec, "
        "ktorý hovorí, či pôvodná hodnota chýbala. Model si tak môže "
        "samostatne naučiť vplyv samotného faktu chýbania, čo je dôležité "
        "najmä pri MNAR (kde chýbanie nesie informáciu).",

        "none — implicitné spracovanie NaN modelom: chýbajúce hodnoty "
        "sa neimputujú a predávajú sa priamo modelu. TabPFN, TabICL, "
        "CatBoost, XGBoost a LightGBM dokážu pracovať s NaN natívne; "
        "klasické modely, ktoré to nepodporujú (LR, RF, GB, SVM, MLP), "
        "boli pri možnosti none preskočené.",
    ])
    add_paragraph(
        doc,
        "Tento výber pokrýva tri kvalitatívne odlišné prístupy: "
        "univariantné štatistické imputácie (mean, median), modelové "
        "imputácie využívajúce ostatné premenné (kNN, MICE, "
        "mice_indicator) a stratégiu, ktorá sa imputácii úplne vyhýba "
        "(none). Vďaka tomu môžeme posúdiť nielen, ktorý imputer je "
        "najlepší, ale aj či je explicitná imputácia vôbec potrebná.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "3.3 Modely", level=2)
    add_paragraph(
        doc,
        "Boli porovnávané dve skupiny modelov. Skupina foundation "
        "obsahuje pretrénované modely TabPFN [1] a TabICL [2]. Skupina "
        "klasických modelov obsahuje Logistic Regression, Random Forest, "
        "Gradient Boosting, XGBoost, LightGBM, SVM, MLP a CatBoost. "
        "CatBoost je v tejto práci klasifikovaný ako klasický NaN-aware "
        "tree ensemble baseline a v agregátnych metrikách Classical vs "
        "Foundation je započítaný do skupiny Classical. Skupina "
        "Foundation tak obsahuje iba TabPFN a TabICL.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Modely boli porovnávané pri použití defaultných alebo vopred "
        "pevne nastavených hyperparametrov. V experimentoch nebolo "
        "vykonané systematické ladenie hyperparametrov pomocou grid search "
        "ani random search; cieľom bolo posúdiť robustnosť modelov pri "
        "rovnakom experimentálnom nastavení, nie maximalizovať výkon "
        "každého modelu individuálnym tuningom. Pre Logistic Regression, "
        "SVM a MLP bola na tréningové dáta aplikovaná štandardizácia "
        "(StandardScaler), ktorá bola fitovaná iba na train sete.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "3.4 Threshold tuning a metriky", level=2)
    add_paragraph(
        doc,
        "Optimalizovaný bol iba rozhodovací prah pri binárnej "
        "klasifikácii. Tréningová množina bola rozdelená na 80 % fit "
        "a 20 % vnútorný validačný split, na ktorom sa hľadal threshold "
        "maximalizujúci balanced accuracy. Model bol následne pretrénovaný "
        "na celej tréningovej množine a vybraný threshold bol aplikovaný "
        "na testovacie predikcie. Test set sa pri tuningu nepoužil, takže "
        "nedochádza k test leakage.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Vzhľadom na výrazne nevyvážené datasety samotná accuracy nestačí "
        "ako hlavná metrika. Hlavnými hodnotenými metrikami sú balanced "
        "accuracy, macro F1, PR-AUC a recall minoritnej triedy. ROC-AUC a "
        "accuracy sú uvádzané iba ako doplnkové metriky.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "4. Výsledky", level=1)

    add_heading(doc, "4.1 Polish Companies 1-year", level=2)
    add_paragraph(
        doc,
        "Priemerná accuracy bola 0,8304, balanced accuracy 0,7865, PR-AUC 0,4899 a "
        "recall minoritnej triedy 0,7390. Najlepším modelom bol jasne "
        "TabPFN s balanced accuracy 0,9169 a PR-AUC 0,8403, nasledovaný "
        "LightGBM (0,8560), XGBoost (0,8483) a TabICL (0,8373). Foundation "
        "modely (TabPFN, TabICL) v priemere prekonali klasické modely:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_bullets(doc, [
        "Classical models: balanced accuracy 0,7614, macro F1 0,5673, "
        "PR-AUC 0,4222.",
        "Foundation models: balanced accuracy 0,8767, macro F1 0,7640, "
        "PR-AUC 0,7328.",
    ])
    add_paragraph(
        doc,
        "Distribučné porovnanie skupín je vizualizované na Obr. 1a "
        "(accuracy, balanced accuracy) a Obr. 1b (macro F1, PR-AUC).",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_split_figure(
        doc, VIZ["polish"] / "classical_vs_foundation.png",
        "Obr. 1a — Polish 1-year: accuracy a balanced accuracy pre "
        "skupiny Classical a Foundation cez všetky experimentálne scenáre.",
        "Obr. 1b — Polish 1-year: macro F1 a PR-AUC pre rovnaké skupiny.",
    )
    add_paragraph(
        doc,
        "Pri rastúcej miere chýbania bol TabPFN najstabilnejší — balanced "
        "accuracy klesla z 0,9452 pri 5 % na 0,8538 pri 40 %. TabICL "
        "klesol z 0,9189 na 0,7252, LightGBM z 0,8968 na 0,7775 a XGBoost "
        "z 0,8944 na 0,7719. Z hľadiska imputácie bola najlepšou "
        "stratégiou none (balanced accuracy 0,8678), nasledovaná "
        "MICE + indicator (0,8124); kNN bola najslabšia (0,7477).",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Tieto trendy naprieč modelmi a úrovňami missingness sumarizuje "
        "teplotná mapa na Obr. 2.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_figure(
        doc, VIZ["polish"] / "stability_heatmap.png",
        "Obr. 2 — Polish 1-year: balanced accuracy pre kombinácie "
        "(model × miera chýbania). Tmavšia zelená znamená vyššiu "
        "stabilitu.",
        width_cm=12.5,
    )

    add_heading(doc, "4.2 Slovak Manufacture 13", level=2)
    add_paragraph(
        doc,
        "Priemerná accuracy bola 0,7710, balanced accuracy 0,7172, PR-AUC 0,2032 a "
        "recall minoritnej triedy 0,6625. Tento dataset si zaslúži "
        "osobitnú pozornosť — testovacia množina obsahuje iba 6 "
        "bankrotujúcich prípadov, takže každá chyba sa silno prejaví v "
        "metrikách. Najlepším modelom bol TabICL s balanced accuracy "
        "0,8353, ktorý prekonal aj TabPFN (0,7966) aj XGBoost (0,7802) a "
        "LightGBM (0,7662). CatBoost dosiahol 0,7418. Skupinové porovnanie:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_bullets(doc, [
        "Classical models: balanced accuracy 0,6896, macro F1 0,4441, "
        "PR-AUC 0,1720.",
        "Foundation models: balanced accuracy 0,8159, macro F1 0,5167, "
        "PR-AUC 0,3149.",
    ])
    add_paragraph(
        doc,
        "Rozdiel medzi skupinami Classical a Foundation je graficky "
        "zachytený na Obr. 3a a Obr. 3b.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_split_figure(
        doc, VIZ["slovak"] / "classical_vs_foundation.png",
        "Obr. 3a — Slovak Manufacture 13: accuracy a balanced accuracy. "
        "Rozdiel medzi skupinami je výrazný napriek extrémnej nevyváženosti.",
        "Obr. 3b — Slovak Manufacture 13: macro F1 a PR-AUC. PR-AUC je "
        "celkovo nízka, čo odráža iba 6 pozitívnych prípadov v teste.",
    )
    add_paragraph(
        doc,
        "Robustnosť TabICL bola podľa mechanizmu pomerne stabilná "
        "(MAR 0,8315, MCAR 0,8249, MNAR 0,8525). Pri zvyšovaní miery "
        "chýbania TabICL kolísal medzi 0,80 a 0,87, kým TabPFN sa "
        "pohyboval medzi 0,75 a 0,83. MLP bol veľmi slabý (balanced "
        "accuracy približne 0,36–0,39). Najlepšou imputačnou stratégiou "
        "bola opäť none (0,8014), s veľkým odstupom pred ostatnými "
        "metódami (kNN 0,7147, MICE 0,7119).",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "4.3 Taiwan Bankruptcy", level=2)
    add_paragraph(
        doc,
        "Priemerná accuracy bola 0,8215, balanced accuracy 0,8406, PR-AUC 0,4438 a recall "
        "minoritnej triedy 0,8609. Najlepším modelom bol TabPFN "
        "(balanced accuracy 0,8741), ale s veľmi tesným XGBoost (0,8711) "
        "a Gradient Boosting (0,8651). TabICL dosiahol 0,8642 a LightGBM "
        "0,8633. CatBoost mal 0,8525. Foundation modely boli mierne lepšie:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_bullets(doc, [
        "Classical models: balanced accuracy 0,8326, macro F1 0,5651, "
        "PR-AUC 0,4141.",
        "Foundation models: balanced accuracy 0,8691, macro F1 0,6149, "
        "PR-AUC 0,5503.",
    ])
    add_paragraph(
        doc,
        "Porovnanie distribúcií metrík pre Taiwan je uvedené na Obr. 4a "
        "a Obr. 4b.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_split_figure(
        doc, VIZ["taiwan"] / "classical_vs_foundation.png",
        "Obr. 4a — Taiwan Bankruptcy: accuracy a balanced accuracy. "
        "Rozdiel medzi skupinami je menší ako pri Polish a Slovak.",
        "Obr. 4b — Taiwan Bankruptcy: macro F1 a PR-AUC. Klasické "
        "boosting modely (XGBoost, LightGBM) sú veľmi konkurencieschopné.",
    )
    add_paragraph(
        doc,
        "Stabilita TabPFN pri rastúcej miere chýbania bola výborná — "
        "balanced accuracy klesla iba z 0,8843 (5 %) na 0,8592 (40 %). "
        "TabICL bol takmer plochý (0,8722 → 0,8618). XGBoost si držal "
        "vysokú úroveň približne 0,87–0,88. Random Forest sa pohyboval "
        "stabilne v rozsahu 0,855–0,861. MLP bol slabší a klesal "
        "rýchlejšie. Najlepšou imputačnou stratégiou bola opäť none "
        "(0,8691), nasledovaná MICE + indicator (0,8416). Z hľadiska "
        "tréningového času zostáva XGBoost a LightGBM najrýchlejšie; "
        "TabPFN je výrazne pomalší kvôli inferencii pretrénovaného "
        "modelu, čo treba zohľadniť pri produkčnom nasadení.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Závislosť výkonu od miery chýbania v scenári MCAR je priamo "
        "znázornená na Obr. 5a a Obr. 5b.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_split_figure(
        doc, VIZ["taiwan"] / "missing_rate_MCAR.png",
        "Obr. 5a — Taiwan Bankruptcy, MCAR: accuracy a balanced "
        "accuracy v závislosti od miery chýbania. Foundation modely "
        "(plné čiary) sa degradujú pomalšie.",
        "Obr. 5b — Taiwan Bankruptcy, MCAR: macro F1 a PR-AUC pri "
        "tých istých scenároch.",
    )

    add_heading(doc, "4.4 Porovnanie metód spracovania chýbajúcich hodnôt", level=2)
    add_paragraph(
        doc,
        "Ak agregujeme balanced accuracy cez všetky modely, mechanizmy a "
        "miery chýbania, dostaneme jednotné poradie metód imputácie pre "
        "každý dataset (Tabuľka 2). Vo všetkých troch prípadoch je "
        "najlepšia stratégia none, čo znamená, že implicitné spracovanie "
        "NaN modelom dosahuje vyššiu balanced accuracy ako akákoľvek "
        "explicitná imputácia. Tento výsledok je výrazný — na Polish "
        "datasete je rozdiel medzi none (0,8678) a druhou najlepšou "
        "možnosťou MICE + indicator (0,8124) viac ako 5 percentuálnych "
        "bodov.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    make_table(
        doc,
        ["Imputácia", "Polish", "Slovak", "Taiwan"],
        [
            ["none",            "0,8678", "0,8014", "0,8691"],
            ["mice_indicator",  "0,8124", "0,7049", "0,8416"],
            ["mean",            "0,7843", "0,7083", "0,8339"],
            ["MICE",            "0,7809", "0,7119", "0,8367"],
            ["median",          "0,7672", "0,7039", "0,8394"],
            ["kNN",             "0,7477", "0,7147", "0,8370"],
        ],
        widths_cm=[3.6, 3.0, 3.0, 3.0],
    )
    add_paragraph(
        doc,
        "Tab. 2 — Priemerná balanced accuracy podľa metódy spracovania "
        "chýbajúcich hodnôt, agregovaná cez všetky modely, mechanizmy a "
        "miery chýbania. Tučné hodnoty by mali zvýrazniť najlepšiu "
        "stratégiu pre každý dataset (vždy none).",
        italic=True, size=9,
    )
    add_paragraph(
        doc,
        "Pri porovnaní medzi imputačnými metódami pozorujeme niekoľko "
        "stabilných vzorcov. Po stratégii none je druhá najlepšia "
        "spravidla mice_indicator (na Polish a Taiwan), čo potvrdzuje, "
        "že informácia o samotnom fakte chýbania má pre model praktickú "
        "hodnotu. Jednoduché štatistické metódy mean a median sú "
        "konkurencieschopné len pri nižších mierach chýbania a na "
        "robustnejších datasetoch (Taiwan), kým pri vyššej miere a "
        "menších datasetoch (Slovak) prepadávajú. kNN imputácia bola na "
        "Polish dokonca najslabšia, čo môže byť spôsobené tým, že "
        "metrika podobnosti vo vysokorozmernom finančnom priestore nie "
        "je dostatočne diskriminatívna.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Výpočtová náročnosť sa medzi imputermi líši rádovo. Na základe "
        "imputation benchmark tabuliek sú MICE a MICE + indicator v "
        "priemere približne 6,2× pomalšie ako kNN a približne 790–936× "
        "pomalšie ako jednoduché mean/median imputácie (v porovnaní s "
        "pass-through stratégiou none ide o viac než 1600×). V praxi to "
        "znamená, že MICE metódy často dominujú celkovému runtime "
        "experimentu bez zodpovedajúceho zisku v balanced accuracy.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Pri IterativeImputer (MICE) bol v konfigurácii použitý "
        "max_iter = 30. Na náročnejších scenároch sa opakovane objavilo, "
        "že early stopping kritérium nebolo splnené a imputácia skončila "
        "na limite počtu iterácií (napr. Slovak a Taiwan pri kontrolnom "
        "teste dosiahli n_iter_ = 30 z 30). Preto je vhodné interpretovať "
        "MICE výsledky ako kompromis medzi kvalitou a vysokou časovou "
        "cenou pri nedokonvergovaných behoch.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Z výskumnej otázky o konkurencieschopnosti implicitného "
        "spracovania NaN voči explicitnej imputácii teda vyplýva jasná "
        "odpoveď: pre modely, ktoré vedia s chýbajúcimi hodnotami "
        "pracovať natívne (TabPFN, TabICL, CatBoost, XGBoost, LightGBM), "
        "je none nielen konkurencieschopné, ale spravidla najlepšie. "
        "Explicitnú imputáciu má zmysel používať primárne tam, kde "
        "model NaN nepodporuje, a v takom prípade je vhodné kombinovať "
        "MICE s indikátormi chýbania.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "5. Diskusia a obmedzenia", level=1)
    add_paragraph(
        doc,
        "Vo všetkých troch datasetoch foundation modely prekonali alebo "
        "vyrovnali klasické modely v balanced accuracy aj PR-AUC. TabPFN "
        "bol najlepší na Polish a Taiwan, kým TabICL bol najsilnejší na "
        "Slovak Manufacture 13 — práve na najmenšom a najnerovnomernejšom "
        "datasete. To naznačuje, že TabICL môže byť výhodný pri veľmi "
        "malých minoritných triedach, kde má TabPFN menej kontextu. "
        "Najpozoruhodnejším pozorovaním je, že na všetkých troch "
        "datasetoch bola najlepšou stratégiou ponechanie NaN hodnôt "
        "modelu (none) — implicitné spracovanie chýbajúcich hodnôt teda "
        "bolo konkurencieschopné a často lepšie ako explicitná imputácia.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Praktické obmedzenia experimentu sú nasledovné. TabPFN má limity "
        "na veľkosť datasetu a počet príznakov a v niektorých "
        "konfiguráciách vyžaduje licenčný token a GPU. TabICL má vlastné "
        "implementačné obmedzenia. Slovak Manufacture 13 obsahuje len 30 "
        "bankrotujúcich firiem, čo robí jeho výsledky štatisticky veľmi "
        "citlivými. Účinok chýbajúcich hodnôt bol simulovaný kontrolovane, "
        "čo neodráža plne všetky reálne scenáre prirodzeného chýbania. "
        "Modely boli porovnávané pri defaultných hyperparametroch; "
        "individuálny tuning by mohol klasickým modelom (najmä MLP a SVM) "
        "pomôcť, ale cieľom bolo objektívne porovnanie pri rovnakom nastavení.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "6. Záver", level=1)
    add_paragraph(
        doc,
        "Práca experimentálne potvrdila, že pretrénované tabular "
        "foundation modely TabPFN a TabICL sú prakticky použiteľné pri "
        "neúplných tabuľkových dátach. V priemere prekonali všetky "
        "porovnávané klasické modely v metrikách relevantných pre "
        "nevyváženú klasifikáciu, pričom najlepšou stratégiou spracovania "
        "chýbajúcich hodnôt bolo ich implicitné spracovanie modelom. "
        "Klasické boosting modely (XGBoost, LightGBM, CatBoost) "
        "zostávajú silnými baseline modelmi, najmä pri väčších datasetoch "
        "alebo v prostrediach bez GPU. Pre malé a silne nevyvážené úlohy "
        "(napr. Slovak Manufacture 13) odporúčame TabICL, pre stredne "
        "veľké datasety s implicitným spracovaním NaN TabPFN.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "Literatúra", level=1)
    refs = [
        "[1] N. Hollmann, S. Müller, L. Purucker, A. Krishnakumar, M. "
        "Körfer, S. B. Hoo, R. T. Schirrmeister, and F. Hutter, "
        "\"Accurate predictions on small data with a tabular foundation "
        "model,\" Nature, vol. 637, pp. 319–326, 2025. [Online]. "
        "Available: https://www.nature.com/articles/s41586-024-08328-6",

        "[2] J. Qu, D. Holzmüller, G. Varoquaux, and M. Le Morvan, "
        "\"TabICL: A Tabular Foundation Model for In-Context Learning on "
        "Large Data,\" arXiv preprint arXiv:2502.05564, 2025. [Online]. "
        "Available: https://arxiv.org/abs/2502.05564",

        "[3] J. Korol, E. Šáteková, and Z. Kostiviarová, \"An "
        "experimental survey of imbalanced learning algorithms for "
        "bankruptcy prediction,\" Artificial Intelligence Review, 2025. "
        "[Online]. Available: "
        "https://link.springer.com/article/10.1007/s10462-025-11107-y",

        "[4] D. B. Rubin, \"Inference and missing data,\" Biometrika, "
        "vol. 63, no. 3, pp. 581–592, 1976.",

        "[5] S. van Buuren and K. Groothuis-Oudshoorn, \"mice: Multivariate "
        "Imputation by Chained Equations in R,\" Journal of Statistical "
        "Software, vol. 45, no. 3, pp. 1–67, 2011.",

        "[6] T. Chen and C. Guestrin, \"XGBoost: A Scalable Tree Boosting "
        "System,\" in Proc. 22nd ACM SIGKDD Int. Conf. on Knowledge "
        "Discovery and Data Mining, 2016, pp. 785–794.",

        "[7] G. Ke et al., \"LightGBM: A Highly Efficient Gradient "
        "Boosting Decision Tree,\" in Advances in Neural Information "
        "Processing Systems, vol. 30, 2017, pp. 3146–3154.",

        "[8] L. Prokhorenkova, G. Gusev, A. Vorobev, A. V. Dorogush, and "
        "A. Gulin, \"CatBoost: unbiased boosting with categorical "
        "features,\" in Advances in Neural Information Processing "
        "Systems, vol. 31, 2018.",

        "[9] F. Pedregosa et al., \"Scikit-learn: Machine Learning in "
        "Python,\" Journal of Machine Learning Research, vol. 12, "
        "pp. 2825–2830, 2011.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        run = p.add_run(r)
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)

    target = OUT_PATH
    try:
        doc.save(target)
    except PermissionError:
        target = OUT_PATH.with_name(OUT_PATH.stem + "_v2.docx")
        doc.save(target)
        print("(Pôvodný súbor je otvorený v aplikácii, ukladám do v2.)")
    print(f"Saved: {target}")
    print(f"Size: {target.stat().st_size} bytes")


if __name__ == "__main__":
    build()
